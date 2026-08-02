from __future__ import annotations

import argparse
import hashlib
import re
import shutil
import tempfile
import uuid
import zipfile
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.document import Document as _Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Mm, Pt, RGBColor, Twips
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
LOGO = ROOT / "assets" / "logo" / "png" / "INTAG_Logo_Primary_Horizontal_RGB_v1_1280w.png"
SYMBOL_SVG = ROOT / "assets" / "logo" / "svg" / "INTAG_Symbol_Color_RGB_v1.svg"
FONT_DIR = ROOT / "assets" / "fonts" / "IBMPlexSansArabic"
FONT_REGULAR = FONT_DIR / "IBMPlexSansArabic-Regular.ttf"
FONT_BOLD = FONT_DIR / "IBMPlexSansArabic-Bold.ttf"
POPPINS_FONT_DIR = ROOT / "assets" / "fonts" / "Poppins"
POPPINS_REGULAR = POPPINS_FONT_DIR / "Poppins-Regular.ttf"
POPPINS_BOLD = POPPINS_FONT_DIR / "Poppins-Bold.ttf"
PROPOSAL_OUT = ROOT / "templates" / "proposal" / "INTAG_Proposal_Starter_AR_v2.docx"
PROPOSAL_EN_OUT = ROOT / "templates" / "proposal" / "INTAG_Proposal_Starter_EN_v2.docx"
LETTERHEAD_OUT = ROOT / "templates" / "stationery" / "INTAG_Letterhead_Template_AR_v2.docx"
PROPOSAL_AR_REFERENCE_SHA256 = "f1c4c1fab0f4d46d3d83c63f8351cc0b87a5337d95155a826f3b5538c73b25d8"


# INTAG v2.0 brand tokens. Logo 02 is approved; the remaining identity
# decisions stay under review until their named owners approve them.
INK = "0B0D10"
MIDNIGHT = "071B2D"
BLUE = "2369B1"
BLUE_DEEP = "134B82"
GREEN = "2AB687"
GREEN_DEEP = "087058"
MINERAL = "F4F1E9"
MIST = "E8F1F2"
ALUMINUM = "C8CED1"
SLATE = "50616F"
WHITE = "FFFFFF"

DISPLAY_FONT = "IBM Plex Sans Arabic"
BODY_FONT = "IBM Plex Sans Arabic"
ARABIC_LANG = "ar-EG"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, color=ALUMINUM, size=6, edges=("top", "start", "bottom", "end")) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in edges:
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_paragraph_border(paragraph, *, top=None, bottom=None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge, color in (("top", top), ("bottom", bottom)):
        if color is None:
            continue
        node = OxmlElement("w:" + edge)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "8")
        node.set(qn("w:space"), "6")
        node.set(qn("w:color"), color)
        p_bdr.append(node)


def set_run_font(run, *, name=BODY_FONT, size=None, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn("w:" + key), name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), ARABIC_LANG)
    lang.set(qn("w:bidi"), ARABIC_LANG)
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")
    if size is not None:
        run.font.size = Pt(size)
        sz_cs = r_pr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            r_pr.append(sz_cs)
        sz_cs.set(qn("w:val"), str(int(round(size * 2))))
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
        b_cs = r_pr.find(qn("w:bCs"))
        if b_cs is None:
            b_cs = OxmlElement("w:bCs")
            r_pr.append(b_cs)
        b_cs.set(qn("w:val"), "1" if bold else "0")
    if italic is not None:
        run.italic = italic
        i_cs = r_pr.find(qn("w:iCs"))
        if i_cs is None:
            i_cs = OxmlElement("w:iCs")
            r_pr.append(i_cs)
        i_cs.set(qn("w:val"), "1" if italic else "0")


def set_paragraph_rtl(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    paragraph.alignment = alignment
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        jc.set(qn("w:val"), "start")
    elif alignment == WD_ALIGN_PARAGRAPH.LEFT:
        jc.set(qn("w:val"), "end")
    elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        jc.set(qn("w:val"), "both")
    else:
        jc.set(qn("w:val"), "center")


def set_keep(paragraph, *, keep_with_next=False, keep_together=False, widow_control=True) -> None:
    fmt = paragraph.paragraph_format
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = keep_together
    fmt.widow_control = widow_control


def add_alt_text(inline_shape, description: str, title: str = "شعار إنتاغ") -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", title)


def add_field(paragraph, instruction: str, display_text: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " " + instruction + " "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run_font(run, name=BODY_FONT, size=8, color=SLATE)


def configure_a4_section(section, *, top=0.72, bottom=0.68) -> int:
    # Named override: INTAG A4 RTL. It preserves the preset's exact 9360 DXA
    # content width while using the locally appropriate A4 page size.
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    content_width = 9360
    side_space = int(section.page_width.twips) - content_width
    section.left_margin = Twips(side_space // 2)
    section.right_margin = Twips(side_space - side_space // 2)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.30)
    return content_width


def set_style_font(style, *, name=BODY_FONT, size=10.5, color=INK, bold=None, italic=None) -> None:
    style.font.name = name
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:cs"), name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), ARABIC_LANG)
    lang.set(qn("w:bidi"), ARABIC_LANG)
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")
    style.font.size = Pt(size)
    sz_cs = r_pr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        r_pr.append(sz_cs)
    sz_cs.set(qn("w:val"), str(int(round(size * 2))))
    style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold
        b_cs = r_pr.find(qn("w:bCs"))
        if b_cs is None:
            b_cs = OxmlElement("w:bCs")
            r_pr.append(b_cs)
        b_cs.set(qn("w:val"), "1" if bold else "0")
    if italic is not None:
        style.font.italic = italic
        i_cs = r_pr.find(qn("w:iCs"))
        if i_cs is None:
            i_cs = OxmlElement("w:iCs")
            r_pr.append(i_cs)
        i_cs.set(qn("w:val"), "1" if italic else "0")


def set_style_rtl(style, *, alignment=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    p_pr = style._element.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    style.paragraph_format.alignment = alignment
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        jc.set(qn("w:val"), "start")
    elif alignment == WD_ALIGN_PARAGRAPH.LEFT:
        jc.set(qn("w:val"), "end")
    elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        jc.set(qn("w:val"), "both")
    else:
        jc.set(qn("w:val"), "center")


def style_or_add(doc: _Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def configure_styles(doc: _Document, *, preset: str) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    set_style_rtl(
        normal,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY if preset == "narrative_proposal" else WD_ALIGN_PARAGRAPH.RIGHT,
    )
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8 if preset == "narrative_proposal" else 6)
    normal.paragraph_format.line_spacing = 1.333 if preset == "narrative_proposal" else 1.10
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    set_style_font(title, name=DISPLAY_FONT, size=30 if preset == "narrative_proposal" else 26, color=MIDNIGHT, bold=True)
    set_style_rtl(title)
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.0
    title.paragraph_format.keep_with_next = True

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, name=DISPLAY_FONT, size=13, color=SLATE, italic=False)
    set_style_rtl(subtitle)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.15
    subtitle.paragraph_format.keep_with_next = True

    heading_values = {
        # Named Arabic legibility override: slightly larger headings than the
        # base preset while preserving its spacing rhythm and color system.
        "Heading 1": (18, BLUE_DEEP, 18 if preset == "narrative_proposal" else 16, 10 if preset == "narrative_proposal" else 8),
        "Heading 2": (14, BLUE_DEEP, 12, 6),
        "Heading 3": (12.5, GREEN_DEEP, 8, 4),
    }
    for name, (size, color, before, after) in heading_values.items():
        style = doc.styles[name]
        set_style_font(style, name=DISPLAY_FONT, size=size, color=color, bold=True)
        set_style_rtl(style)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.05
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    kicker = style_or_add(doc, "INTAG Kicker")
    set_style_font(kicker, name=DISPLAY_FONT, size=8.5, color=GREEN_DEEP, bold=True)
    set_style_rtl(kicker)
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(7)
    kicker.paragraph_format.keep_with_next = True

    lead = style_or_add(doc, "INTAG Lead")
    set_style_font(lead, name=DISPLAY_FONT, size=13, color=MIDNIGHT, bold=True)
    set_style_rtl(lead)
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.20
    lead.paragraph_format.keep_together = True

    meta = style_or_add(doc, "INTAG Meta")
    set_style_font(meta, name=BODY_FONT, size=8.5, color=SLATE)
    set_style_rtl(meta)
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(3)
    meta.paragraph_format.line_spacing = 1.0

    note = style_or_add(doc, "INTAG Note")
    set_style_font(note, name=BODY_FONT, size=9, color=MIDNIGHT)
    set_style_rtl(note)
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.15

    placeholder = style_or_add(doc, "INTAG Placeholder")
    set_style_font(placeholder, name=BODY_FONT, size=10.5, color=SLATE, italic=False)
    set_style_rtl(placeholder)
    placeholder.paragraph_format.space_before = Pt(0)
    placeholder.paragraph_format.space_after = Pt(8)
    placeholder.paragraph_format.line_spacing = 1.25

    bullet = style_or_add(doc, "INTAG Bullet")
    set_style_font(bullet, size=10.75, color=INK)
    set_style_rtl(bullet)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(4 if preset == "narrative_proposal" else 8)
    bullet.paragraph_format.line_spacing = 1.208 if preset == "narrative_proposal" else 1.167

    number = style_or_add(doc, "INTAG Number")
    set_style_font(number, size=10.75, color=INK)
    set_style_rtl(number)
    number.paragraph_format.space_before = Pt(0)
    number.paragraph_format.space_after = Pt(4 if preset == "narrative_proposal" else 8)
    number.paragraph_format.line_spacing = 1.208 if preset == "narrative_proposal" else 1.167

    table_header = style_or_add(doc, "INTAG Table Header")
    set_style_font(table_header, size=9.25, color=WHITE, bold=True)
    set_style_rtl(table_header)
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(0)
    table_header.paragraph_format.line_spacing = 1.05

    table_body = style_or_add(doc, "INTAG Table Body")
    set_style_font(table_body, size=9.25, color=INK)
    set_style_rtl(table_body)
    table_body.paragraph_format.space_before = Pt(0)
    table_body.paragraph_format.space_after = Pt(0)
    table_body.paragraph_format.line_spacing = 1.15

    signature = style_or_add(doc, "INTAG Signature")
    set_style_font(signature, size=9.25, color=INK)
    set_style_rtl(signature)
    signature.paragraph_format.space_before = Pt(0)
    signature.paragraph_format.space_after = Pt(5)
    signature.paragraph_format.line_spacing = 1.05


def install_numbering(doc: _Document, *, preset: str) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abstract = (max(abstract_ids) + 1) if abstract_ids else 1
    next_num = (max(num_ids) + 1) if num_ids else 1

    if preset == "narrative_proposal":
        marker, left, hanging, line, after = 260, 540, 280, 290, 80
    else:
        marker, left, hanging, line, after = 360, 720, 360, 280, 160

    def build_abstract(abstract_id: int, fmt: str, text: str):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "right")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:right"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), str(after))
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        if fmt == "bullet":
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), BODY_FONT)
            r_fonts.set(qn("w:hAnsi"), BODY_FONT)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        return abstract

    def build_num(num_id: int, abstract_id: int):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        return num

    # OOXML requires all abstractNum elements before num instances. Keeping this
    # order prevents Word from silently remapping the bullet list as decimal.
    numbering.append(build_abstract(next_abstract, "bullet", "•"))
    numbering.append(build_abstract(next_abstract + 1, "decimal", "%1."))
    numbering.append(build_num(next_num, next_abstract))
    numbering.append(build_num(next_num + 1, next_abstract + 1))
    return next_num, next_num + 1


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))


def add_bullet(doc_or_cell, text: str, bullet_id: int):
    if hasattr(doc_or_cell, "_tc") and len(doc_or_cell.paragraphs) == 1 and not doc_or_cell.paragraphs[0].text:
        p = doc_or_cell.paragraphs[0]
        p.style = "INTAG Bullet"
    else:
        p = doc_or_cell.add_paragraph(style="INTAG Bullet")
    apply_numbering(p, bullet_id)
    set_paragraph_rtl(p)
    run = p.add_run(text)
    set_run_font(run, size=10.75, color=INK)
    set_keep(p, keep_together=True)
    return p


def add_number(doc_or_cell, text: str, number_id: int):
    if hasattr(doc_or_cell, "_tc") and len(doc_or_cell.paragraphs) == 1 and not doc_or_cell.paragraphs[0].text:
        p = doc_or_cell.paragraphs[0]
        p.style = "INTAG Number"
    else:
        p = doc_or_cell.add_paragraph(style="INTAG Number")
    apply_numbering(p, number_id)
    set_paragraph_rtl(p)
    run = p.add_run(text)
    set_run_font(run, size=10.75, color=INK)
    set_keep(p, keep_together=True)
    return p


def set_table_geometry(table, widths_twips: Sequence[int], *, indent_twips=120) -> None:
    total = sum(widths_twips)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    tbl_pr = table._tbl.tblPr
    bidi_visual = tbl_pr.find(qn("w:bidiVisual"))
    if bidi_visual is None:
        bidi_visual = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi_visual)
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_twips[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = int(widths_twips[idx] * 635)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def clear_cell(cell) -> None:
    p = cell.paragraphs[0]
    for run in list(p.runs):
        p._p.remove(run._r)


def add_cell_text(cell, text: str, *, style="INTAG Table Body", color=INK, bold=False, align=None) -> None:
    p = cell.paragraphs[0]
    p.style = style
    set_paragraph_rtl(p, alignment=align or WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(text)
    set_run_font(run, size=9.25, color=color, bold=bold)
    set_keep(p, keep_together=True)


def add_data_table(doc: _Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths_twips: Sequence[int], *, status_column=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths_twips)
    mark_header_row(table.rows[0])
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        clear_cell(cell)
        set_cell_shading(cell, MIDNIGHT)
        set_cell_border(cell, color=MIDNIGHT, size=6)
        add_cell_text(cell, header, style="INTAG Table Header", color=WHITE, bold=True)
    for i, row_values in enumerate(rows, start=1):
        for j, value in enumerate(row_values):
            cell = table.rows[i].cells[j]
            clear_cell(cell)
            set_cell_shading(cell, WHITE if i % 2 else "F8FAFA")
            set_cell_border(cell, color=ALUMINUM, size=4)
            align = WD_ALIGN_PARAGRAPH.CENTER if status_column is not None and j == status_column else WD_ALIGN_PARAGRAPH.RIGHT
            add_cell_text(cell, value, color=GREEN_DEEP if status_column is not None and j == status_column else INK, bold=(j == 0), align=align)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc: _Document, text: str, content_width: int, *, label=None, fill=MINERAL, accent=GREEN_DEEP) -> None:
    # Paragraph callout (not a layout table): semantic text with brand shading
    # and a single start-edge accent, plus explicit internal padding.
    p = doc.add_paragraph()
    p.style = "INTAG Note"
    set_paragraph_rtl(p)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "12")
    start.set(qn("w:space"), "8")
    start.set(qn("w:color"), accent)
    borders.append(start)
    p_pr.append(borders)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "120")
    ind.set(qn("w:right"), "120")
    p_pr.append(ind)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    if label:
        r = p.add_run(label + "  ")
        set_run_font(r, size=9, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MIDNIGHT)
    set_keep(p, keep_together=True)


def add_spacer(doc: _Document, points: float) -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(" ")
    set_run_font(r, size=1, color=WHITE)


def set_header_footer(section, *, document_label: str, status_label: str, content_width: int, logo=False) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    p = header.paragraphs[0]
    p.style = "Normal"
    set_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1
    if logo:
        run = p.add_run()
        shape = run.add_picture(str(LOGO), width=Inches(2.05))
        add_alt_text(shape, "شعار إنتاغ للحلول الرقمية، الشعار رقم 02")
    else:
        r = p.add_run("إنتاغ للحلول الرقمية")
        set_run_font(r, name=DISPLAY_FONT, size=8.5, color=BLUE_DEEP, bold=True)

    status = header.add_paragraph()
    status.style = "INTAG Meta"
    set_paragraph_rtl(status)
    status.paragraph_format.space_before = Pt(0)
    status.paragraph_format.space_after = Pt(0)
    status.paragraph_format.line_spacing = 1
    r = status.add_run(status_label)
    set_run_font(r, name=BODY_FONT, size=7.5, color=SLATE, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.style = "Normal"
    set_paragraph_rtl(fp, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.line_spacing = 1
    set_paragraph_border(fp, top=ALUMINUM)
    r = fp.add_run(document_label + "  ·  قالب عمل قابل للتحرير  ·  الإصدار 2.0  ·  صفحة ")
    set_run_font(r, size=7.5, color=SLATE)
    add_field(fp, "PAGE", "1")
    r = fp.add_run(" من ")
    set_run_font(r, size=7.5, color=SLATE)
    add_field(fp, "NUMPAGES", "1")


def add_doc_properties(doc: _Document, *, title: str, subject: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "إنتاغ للحلول الرقمية"
    props.last_modified_by = "إنتاغ للحلول الرقمية"
    props.keywords = "إنتاغ، قالب عمل، هوية بصرية، مقترح"
    props.comments = "قالب عمل. يجب استبدال جميع الحقول المحاطة بأقواس مربعة قبل الاستخدام."


def verify_logo_geometry() -> None:
    if not SYMBOL_SVG.exists():
        raise FileNotFoundError(SYMBOL_SVG)
    source = SYMBOL_SVG.read_text(encoding="utf-8")
    path_data = re.findall(r'<path\s+d="([^"]+)"', source)
    expected = ["M20 84V28H72", "M108 44V100H56"]
    if path_data != expected:
        raise RuntimeError(
            "Logo 02 geometry gate failed. Expected the two approved orthogonal paths exactly: "
            + " | ".join(expected)
        )
    if any(command in "".join(path_data) for command in "LlCcQqSsTtAa"):
        raise RuntimeError("Logo 02 geometry gate failed: diagonal or curved commands were found.")


def _obfuscate_font(font_path: Path, font_key: uuid.UUID) -> bytes:
    data = bytearray(font_path.read_bytes())
    key = font_key.bytes_le
    for index in range(min(32, len(data))):
        data[index] ^= key[index % 16]
    return bytes(data)


def embed_arabic_fonts(docx_path: Path) -> None:
    """Embed the OFL-licensed regular and bold brand fonts into the DOCX."""
    for font_path in (FONT_REGULAR, FONT_BOLD):
        if not font_path.exists():
            raise FileNotFoundError(font_path)

    with zipfile.ZipFile(docx_path, "r") as source_zip:
        entries = {item.filename: source_zip.read(item.filename) for item in source_zip.infolist()}

    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    font_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font"

    font_table_name = "word/fontTable.xml"
    font_table = etree.fromstring(entries[font_table_name])
    font_node = font_table.find(f"{{{w_ns}}}font[@{{{w_ns}}}name='{BODY_FONT}']")
    if font_node is None:
        font_node = etree.SubElement(font_table, f"{{{w_ns}}}font")
        font_node.set(f"{{{w_ns}}}name", BODY_FONT)
        family = etree.SubElement(font_node, f"{{{w_ns}}}family")
        family.set(f"{{{w_ns}}}val", "swiss")
        charset = etree.SubElement(font_node, f"{{{w_ns}}}charset")
        charset.set(f"{{{w_ns}}}val", "00")
    for tag in ("embedRegular", "embedBold", "embedItalic", "embedBoldItalic"):
        child = font_node.find(f"{{{w_ns}}}{tag}")
        if child is not None:
            font_node.remove(child)

    rels_name = "word/_rels/fontTable.xml.rels"
    if rels_name in entries:
        rels = etree.fromstring(entries[rels_name])
    else:
        rels = etree.Element(f"{{{rel_ns}}}Relationships", nsmap={None: rel_ns})

    specs = [
        ("embedRegular", "rIdIntagArabicRegular", FONT_REGULAR, "IBMPlexSansArabic-Regular.odttf"),
        ("embedBold", "rIdIntagArabicBold", FONT_BOLD, "IBMPlexSansArabic-Bold.odttf"),
    ]
    for _, relation_id, _, _ in specs:
        for relation in list(rels):
            if relation.get("Id") == relation_id:
                rels.remove(relation)

    for tag, relation_id, font_path, target_name in specs:
        font_key = uuid.uuid5(uuid.NAMESPACE_URL, f"intag-v2:{font_path.name}")
        embed = etree.SubElement(font_node, f"{{{w_ns}}}{tag}")
        embed.set(f"{{{r_ns}}}id", relation_id)
        embed.set(f"{{{w_ns}}}fontKey", "{" + str(font_key).upper() + "}")
        embed.set(f"{{{w_ns}}}subsetted", "0")

        relation = etree.SubElement(rels, f"{{{rel_ns}}}Relationship")
        relation.set("Id", relation_id)
        relation.set("Type", font_rel_type)
        relation.set("Target", f"fonts/{target_name}")
        entries[f"word/fonts/{target_name}"] = _obfuscate_font(font_path, font_key)

    content_types = etree.fromstring(entries["[Content_Types].xml"])
    if content_types.find(f"{{{ct_ns}}}Default[@Extension='odttf']") is None:
        default = etree.SubElement(content_types, f"{{{ct_ns}}}Default")
        default.set("Extension", "odttf")
        default.set("ContentType", "application/vnd.openxmlformats-officedocument.obfuscatedFont")

    entries[font_table_name] = etree.tostring(font_table, xml_declaration=True, encoding="UTF-8", standalone="yes")
    entries[rels_name] = etree.tostring(rels, xml_declaration=True, encoding="UTF-8", standalone="yes")
    entries["[Content_Types].xml"] = etree.tostring(
        content_types, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    with tempfile.NamedTemporaryFile(dir=docx_path.parent, suffix=".docx", delete=False) as temp_file:
        temp_path = Path(temp_file.name)
    try:
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as output_zip:
            for name, data in entries.items():
                output_zip.writestr(name, data)
        shutil.move(str(temp_path), str(docx_path))
    finally:
        if temp_path.exists():
            temp_path.unlink()


def audit_visible_language_and_terms(docx_path: Path) -> None:
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(docx_path, "r") as archive:
        visible_parts = [
            name
            for name in archive.namelist()
            if name == "word/document.xml" or re.fullmatch(r"word/(header|footer)\d+\.xml", name)
        ]
        visible_text: list[str] = []
        for name in visible_parts:
            root = etree.fromstring(archive.read(name))
            visible_text.extend(node.text or "" for node in root.findall(f".//{{{w_ns}}}t"))
            for node in root.iter():
                for attribute in ("descr", "title"):
                    if node.get(attribute):
                        visible_text.append(node.get(attribute))
        joined_visible = "\n".join(visible_text)
        core = archive.read("docProps/core.xml").decode("utf-8", errors="replace")

    prohibited = (r"\bA" + r"I\b", "Artificial " + "Intelligence", "الذكاء " + "الاصطناعي")
    for pattern in prohibited:
        if re.search(pattern, joined_visible + "\n" + core, flags=re.IGNORECASE):
            raise RuntimeError(f"Disallowed term detected in visible content or metadata: {pattern}")
    if re.search(r"[A-Za-z]", joined_visible):
        latin_matches = sorted(set(re.findall(r"[A-Za-z][A-Za-z ./:&-]*", joined_visible)))
        raise RuntimeError("Visible non-Arabic text detected: " + " | ".join(latin_matches[:8]))


def finalize_document(doc: _Document, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    embed_arabic_fonts(output_path)
    audit_visible_language_and_terms(output_path)


def add_page_break(doc: _Document) -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_text_paragraph(doc_or_cell, text: str, *, style=None, color=INK, bold=False, italic=False, size=None, align=None, after=None):
    p = doc_or_cell.add_paragraph(style=style)
    set_paragraph_rtl(p, alignment=align or WD_ALIGN_PARAGRAPH.RIGHT)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def build_proposal() -> None:
    doc = Document()
    section = doc.sections[0]
    content_width = configure_a4_section(section)
    configure_styles(doc, preset="narrative_proposal")
    bullet_id, number_id = install_numbering(doc, preset="narrative_proposal")
    set_header_footer(
        section,
        document_label="INTAG PROPOSAL STARTER",
        status_label="PROPOSAL / WORKING",
        content_width=content_width,
        logo=False,
    )
    add_doc_properties(
        doc,
        title="INTAG Proposal Starter v1",
        subject="Working proposal template for client engagements",
    )

    # Page 1 — proposal_centerpiece branded override
    add_spacer(doc, 22)
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(22)
    logo_run = p_logo.add_run()
    logo_shape = logo_run.add_picture(str(LOGO), width=Inches(3.00))
    add_alt_text(logo_shape, "INTAG Digital Solutions horizontal logo")

    p = doc.add_paragraph(style="INTAG Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PROPOSAL / WORKING DRAFT")

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[ENGAGEMENT TITLE]")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[ONE CLEAR OUTCOME-LED SUBTITLE]")

    p = doc.add_paragraph(style="INTAG Meta")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared for ")
    set_run_font(r, size=9, color=SLATE)
    r = p.add_run("[CLIENT / ORGANISATION]")
    set_run_font(r, size=9, color=BLUE_DEEP, bold=True)
    p.paragraph_format.space_after = Pt(20)

    half = content_width // 2
    metadata = add_data_table(
        doc,
        ["DOCUMENT", "OWNER"],
        [
            ("Date\n[DATE]", "Proposal owner\n[OWNER]"),
            ("Version\n[VERSION]", "Prepared by\nINTAG Digital Solutions"),
        ],
        [half, content_width - half],
    )
    for row in metadata.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
    add_spacer(doc, 7)
    add_callout(
        doc,
        "Scope, fees, dates, success targets, and external commitments remain Pending until documented approval by the named owners.",
        content_width,
        label="STATUS",
        fill=MINERAL,
        accent=GREEN_DEEP,
    )
    add_spacer(doc, 18)
    p = doc.add_paragraph(style="INTAG Meta")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("We build what growth needs.  |  From potential to production.")
    set_run_font(r, size=8.5, color=SLATE, bold=True)

    add_page_break(doc)

    # Page 2 — executive objective
    p = doc.add_paragraph(style="INTAG Kicker")
    p.add_run("01 / EXECUTIVE OBJECTIVE")
    h1 = doc.add_paragraph("The objective", style="Heading 1")
    set_keep(h1, keep_with_next=True)
    p = doc.add_paragraph(style="INTAG Lead")
    p.add_run("Define the work required to move [CLIENT] from [CURRENT STATE] toward [DESIRED STATE] through a measurable, approved engagement.")

    add_callout(
        doc,
        "This objective is a working hypothesis. Validate it in discovery before it becomes an approved client commitment.",
        content_width,
        label="WORKING ASSUMPTION",
        fill=MIST,
        accent=BLUE_DEEP,
    )

    doc.add_paragraph("Context", style="Heading 2")
    add_text_paragraph(
        doc,
        "[ONE OR TWO SENTENCES DESCRIBING THE CURRENT SITUATION, USER NEED, COMMERCIAL CONTEXT, OR AVAILABLE EVIDENCE.]",
        style="INTAG Placeholder",
    )

    doc.add_paragraph("Recommended direction", style="Heading 2")
    add_text_paragraph(
        doc,
        "[A CONCISE, EVIDENCE-BASED DIRECTION. FRAME IT AS A TESTABLE RECOMMENDATION; DO NOT PROMISE A RESULT.]",
        style="INTAG Placeholder",
    )

    doc.add_paragraph("Evidence state", style="Heading 2")
    evidence = doc.add_table(rows=2, cols=2)
    set_table_geometry(evidence, [half, content_width - half])
    mark_header_row(evidence.rows[0])
    for idx, title in enumerate(("WHAT WE KNOW", "WHAT REMAINS PENDING")):
        c = evidence.rows[0].cells[idx]
        clear_cell(c)
        set_cell_shading(c, MIDNIGHT)
        set_cell_border(c, color=MIDNIGHT)
        add_cell_text(c, title, style="INTAG Table Header", color=WHITE, bold=True)
    for idx in range(2):
        c = evidence.rows[1].cells[idx]
        clear_cell(c)
        set_cell_shading(c, WHITE if idx == 0 else MINERAL)
        set_cell_border(c, color=ALUMINUM)
    add_bullet(evidence.rows[1].cells[0], "[CONFIRMED FACT OR APPROVED DECISION]", bullet_id)
    add_bullet(evidence.rows[1].cells[0], "[AVAILABLE EVIDENCE OR EXISTING ASSET]", bullet_id)
    add_bullet(evidence.rows[1].cells[1], "[AUDIENCE, PRIORITY, OR BASELINE TO CONFIRM]", bullet_id)
    add_bullet(evidence.rows[1].cells[1], "[BUDGET, TIMELINE, OR TECHNICAL CONSTRAINT]", bullet_id)
    add_spacer(doc, 2)

    doc.add_paragraph("Success criteria", style="Heading 2")
    add_bullet(doc, "[MEASURABLE OUTCOME] — baseline [BASELINE], proposed target [TARGET], measurement [METHOD].", bullet_id)
    add_bullet(doc, "[QUALITY OR EXPERIENCE INDICATOR] — proposed acceptance threshold [THRESHOLD].", bullet_id)
    add_text_paragraph(
        doc,
        "Targets remain Pending until baseline data, measurement ownership, and approval are confirmed.",
        style="INTAG Meta",
        color=SLATE,
        italic=True,
    )

    add_page_break(doc)

    # Page 3 — scope and deliverables
    p = doc.add_paragraph(style="INTAG Kicker")
    p.add_run("02 / SCOPE & DELIVERABLES")
    doc.add_paragraph("Scope & deliverables", style="Heading 1")
    add_text_paragraph(
        doc,
        "Use this page to convert the objective into explicit outputs, acceptance evidence, and ownership. Keep only approved workstreams in the final proposal.",
        style="Normal",
    )

    widths = [int(content_width * 0.20), int(content_width * 0.34), int(content_width * 0.31)]
    widths.append(content_width - sum(widths))
    add_data_table(
        doc,
        ["WORKSTREAM", "PROPOSED DELIVERABLE", "ACCEPTANCE EVIDENCE", "STATUS"],
        [
            ("Strategy & Business", "[STRATEGY OUTPUT]", "[APPROVAL OR DECISION RECORD]", "Pending"),
            ("Brand & Media", "[BRAND / MEDIA OUTPUT]", "[REVIEW CRITERIA OR ASSET LIST]", "Pending"),
            ("Technology & Digital Products", "[DIGITAL PRODUCT OUTPUT]", "[TEST, DEMO, OR HANDOFF EVIDENCE]", "Pending"),
            ("Growth & Activation", "[GROWTH EXPERIMENT OR ACTIVATION]", "[MEASURED RESULT OR LEARNING]", "Pending"),
        ],
        widths,
        status_column=3,
    )

    doc.add_paragraph("Scope boundaries", style="Heading 2")
    add_bullet(doc, "Only deliverables explicitly listed in the approved version are in scope.", bullet_id)
    add_bullet(doc, "Production deployment, integrations, third-party licenses, media spend, and ongoing support remain Pending unless named above.", bullet_id)
    add_bullet(doc, "Dates and fees are not committed until the commercial owner and delivery owner approve the same version.", bullet_id)

    doc.add_paragraph("Dependencies & client inputs", style="Heading 2")
    add_bullet(doc, "[ACCESS, DATA, ASSETS, OR APPROVALS REQUIRED FROM CLIENT]", bullet_id)
    add_bullet(doc, "[INTAG OWNER] to confirm the delivery approach and handoff criteria.", bullet_id)
    add_bullet(doc, "[TECHNICAL OWNER / THIRD PARTY] to approve any production architecture or deployment responsibility.", bullet_id)

    add_callout(
        doc,
        "If a dependency changes, update scope, timing, and commercials through a documented change decision before work proceeds.",
        content_width,
        label="CHANGE CONTROL",
        fill=MIST,
        accent=BLUE_DEEP,
    )

    add_page_break(doc)

    # Page 4 — decisions and next steps
    p = doc.add_paragraph(style="INTAG Kicker")
    p.add_run("03 / DECISIONS & NEXT STEPS")
    doc.add_paragraph("Decisions & next steps", style="Heading 1")
    add_text_paragraph(
        doc,
        "This final page makes approval state and ownership visible before any external commitment or kickoff.",
        style="Normal",
    )

    decision_widths = [int(content_width * 0.38), int(content_width * 0.22), int(content_width * 0.20)]
    decision_widths.append(content_width - sum(decision_widths))
    add_data_table(
        doc,
        ["DECISION", "OWNER", "TARGET DATE", "STATE"],
        [
            ("Objective & success criteria", "[OWNER]", "[DATE]", "Pending"),
            ("Scope & acceptance evidence", "[OWNER]", "[DATE]", "Pending"),
            ("Commercial terms & payment", "[COMMERCIAL OWNER]", "[DATE]", "Pending"),
            ("Technical delivery ownership", "[TECHNICAL OWNER]", "[DATE]", "Pending"),
        ],
        decision_widths,
        status_column=3,
    )

    doc.add_paragraph("Next-step sequence", style="Heading 2")
    add_number(doc, "Confirm the client objective, baseline evidence, and named decision owners.", number_id)
    add_number(doc, "Refine the deliverables, dependencies, acceptance evidence, and exclusions.", number_id)
    add_number(doc, "Complete commercial, legal, and technical review against one version.", number_id)
    add_number(doc, "Approve the proposal, then schedule kickoff and handoff with accountable owners.", number_id)

    doc.add_paragraph("Approval gate", style="Heading 2")
    add_callout(
        doc,
        "Approval confirms the stated document version and scope only. It does not guarantee a sales, growth, accuracy, technical, or delivery outcome.",
        content_width,
        label="NO-GUARANTEE BOUNDARY",
        fill=MINERAL,
        accent=GREEN_DEEP,
    )

    sign = doc.add_table(rows=2, cols=2)
    set_table_geometry(sign, [half, content_width - half])
    mark_header_row(sign.rows[0])
    for idx, title in enumerate(("CLIENT APPROVER", "INTAG APPROVER")):
        c = sign.rows[0].cells[idx]
        clear_cell(c)
        set_cell_shading(c, MIDNIGHT)
        set_cell_border(c, color=MIDNIGHT)
        add_cell_text(c, title, style="INTAG Table Header", color=WHITE, bold=True)
    for idx, label in enumerate(("[CLIENT APPROVER]", "[INTAG APPROVER]")):
        c = sign.rows[1].cells[idx]
        clear_cell(c)
        set_cell_shading(c, WHITE)
        set_cell_border(c, color=ALUMINUM)
        for line_index, text in enumerate((f"Name: {label}", "Role: [ROLE]", "Signature: [PENDING]", "Date: [DATE]", "Approval state: Pending")):
            if line_index == 0:
                p = c.paragraphs[0]
                p.style = "INTAG Signature"
            else:
                p = c.add_paragraph(style="INTAG Signature")
            r = p.add_run(text)
            set_run_font(r, size=8.75, color=INK, bold=text.startswith("Name:"))

    PROPOSAL_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(PROPOSAL_OUT)


def build_letterhead() -> None:
    doc = Document()
    section = doc.sections[0]
    content_width = configure_a4_section(section, left=0.78, right=0.78, top=0.98, bottom=0.72)
    configure_styles(doc, preset="standard_business_brief")
    bullet_id, _ = install_numbering(doc, preset="standard_business_brief")
    set_header_footer(
        section,
        document_label="INTAG CORRESPONDENCE",
        status_label="DIGITAL SOLUTIONS / WORKING",
        content_width=content_width,
        logo=True,
    )
    add_doc_properties(
        doc,
        title="INTAG Letterhead Template v1",
        subject="Working A4 letterhead template",
    )

    p = doc.add_paragraph(style="INTAG Kicker")
    p.add_run("CORRESPONDENCE")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(14)

    meta = doc.add_paragraph(style="INTAG Meta")
    meta.paragraph_format.tab_stops.add_tab_stop(Inches(content_width / 1440), WD_TAB_ALIGNMENT.RIGHT)
    r = meta.add_run("[DATE]")
    set_run_font(r, size=9, color=SLATE, bold=True)
    r = meta.add_run("\t[REFERENCE / VERSION]")
    set_run_font(r, size=8.5, color=SLATE)
    meta.paragraph_format.space_after = Pt(16)

    for line, bold in (
        ("[RECIPIENT NAME]", True),
        ("[TITLE]", False),
        ("[ORGANISATION]", False),
        ("[ADDRESS LINE 1]", False),
        ("[CITY / COUNTRY]", False),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        set_run_font(r, size=10, color=INK, bold=bold)

    add_spacer(doc, 8)
    subject = doc.add_paragraph()
    subject.paragraph_format.space_after = Pt(16)
    set_paragraph_border(subject, bottom=GREEN)
    r = subject.add_run("SUBJECT  ")
    set_run_font(r, size=9, color=GREEN_DEEP, bold=True)
    r = subject.add_run("[SUBJECT]")
    set_run_font(r, size=11, color=MIDNIGHT, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Dear [RECIPIENT],")
    set_run_font(r, size=10.5, color=INK)

    p = doc.add_paragraph(style="INTAG Placeholder")
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("[START TYPING HERE. REPLACE THIS PROMPT WITH THE APPROVED MESSAGE, KEEP PARAGRAPHS SHORT, AND VERIFY ALL FACTS, DATES, OWNERS, AND COMMITMENTS BEFORE SENDING.]")
    set_run_font(r, size=10.5, color=SLATE, italic=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("Kind regards,")
    set_run_font(r, size=10.5, color=INK)

    for line, bold, color in (
        ("[NAME SURNAME]", True, MIDNIGHT),
        ("[TITLE]", False, SLATE),
        ("INTAG Digital Solutions", True, BLUE_DEEP),
        ("[EMAIL]  |  [PHONE]  |  [WEBSITE]", False, SLATE),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run_font(r, size=9.5 if not bold else 10, color=color, bold=bold)

    add_spacer(doc, 42)
    add_callout(
        doc,
        "Replace every bracketed placeholder before sending. This stationery remains Working until Brand Owner approval and final legal/contact details are confirmed.",
        content_width,
        label="TEMPLATE NOTE",
        fill=MINERAL,
        accent=GREEN_DEEP,
    )

    LETTERHEAD_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(LETTERHEAD_OUT)


# v2.0 Arabic-first builders. These definitions intentionally replace the
# legacy working builders above while preserving the reusable OOXML helpers.
def build_proposal() -> None:
    doc = Document()
    section = doc.sections[0]
    content_width = configure_a4_section(section, top=0.72, bottom=0.68)
    configure_styles(doc, preset="narrative_proposal")
    bullet_id, number_id = install_numbering(doc, preset="narrative_proposal")
    set_header_footer(
        section,
        document_label="قالب مقترح إنتاغ",
        status_label="الشعار 02 معتمد، وبقية القرارات قيد المراجعة",
        content_width=content_width,
        logo=False,
    )
    add_doc_properties(
        doc,
        title="قالب مقترح إنتاغ - الإصدار 2.0",
        subject="قالب عربي قابل للتحرير لإعداد مقترحات العمل ومراجعتها واعتمادها",
    )

    # الغلاف: نمط proposal_centerpiece مع معالجة عربية تحمل هوية إنتاغ.
    add_spacer(doc, 16)
    p_logo = doc.add_paragraph()
    set_paragraph_rtl(p_logo, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p_logo.paragraph_format.space_after = Pt(20)
    logo_run = p_logo.add_run()
    logo_shape = logo_run.add_picture(str(LOGO), width=Inches(2.80))
    add_alt_text(logo_shape, "شعار إنتاغ للحلول الرقمية، الشعار رقم 02")

    p = doc.add_paragraph(style="INTAG Kicker")
    set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run("مقترح عمل · نسخة قابلة للتحرير")

    p = doc.add_paragraph(style="Title")
    set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run("[عنوان المشروع أو المبادرة]")

    p = doc.add_paragraph(style="Subtitle")
    set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run("[وصف موجز يوضح النتيجة المطلوبة بلغة مباشرة]")

    p = doc.add_paragraph(style="INTAG Meta")
    set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("مقدم إلى: ")
    set_run_font(r, size=9, color=SLATE)
    r = p.add_run("[اسم العميل أو الجهة]")
    set_run_font(r, size=9, color=BLUE_DEEP, bold=True)
    p.paragraph_format.space_after = Pt(18)

    half = content_width // 2
    metadata = add_data_table(
        doc,
        ["بيانات الوثيقة", "المسؤولية"],
        [
            ("التاريخ\n[التاريخ]", "مالك المقترح\n[اسم المالك]"),
            ("الإصدار\n2.0", "الإعداد\nإنتاغ للحلول الرقمية"),
        ],
        [half, content_width - half],
    )
    for row in metadata.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)

    add_spacer(doc, 6)
    add_callout(
        doc,
        "اعتمد الشعار رقم 02 «الإطار المتصل». أما بقية قرارات الهوية والنطاق والتكلفة والجدول الزمني فما تزال قيد المراجعة والاعتماد.",
        content_width,
        label="حالة الوثيقة",
        fill=MINERAL,
        accent=GREEN_DEEP,
    )
    add_spacer(doc, 15)
    p = doc.add_paragraph(style="INTAG Meta")
    set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("يجب استبدال كل حقل محاط بأقواس مربعة قبل مشاركة النسخة النهائية.")
    set_run_font(r, size=8.5, color=SLATE, bold=True)

    add_page_break(doc)

    # الصفحة 2: الملخص التنفيذي.
    p = doc.add_paragraph(style="INTAG Kicker")
    p.add_run("01 / الملخص التنفيذي")
    doc.add_paragraph("الملخص التنفيذي", style="Heading 1")
    p = doc.add_paragraph(style="INTAG Lead")
    p.add_run("يوضح هذا القسم سبب العمل والنتيجة المطلوبة وما يحتاج إلى اعتماد قبل بدء التنفيذ.")

    doc.add_paragraph("السياق", style="Heading 2")
    add_text_paragraph(
        doc,
        "[اكتب في فقرة قصيرة ما يحدث الآن، ومن يتأثر به، ولماذا يحتاج إلى معالجة في هذا التوقيت.]",
        style="INTAG Placeholder",
    )

    doc.add_paragraph("التحدي والنتيجة المطلوبة", style="Heading 2")
    add_text_paragraph(
        doc,
        "[صغ التحدي من منظور العمل، ثم حدد النتيجة المرغوبة من دون تحويلها إلى وعد أو نتيجة مضمونة.]",
        style="INTAG Placeholder",
    )

    doc.add_paragraph("حالة المعلومات", style="Heading 2")
    add_data_table(
        doc,
        ["التصنيف", "ما يكتب هنا"],
        [
            ("معلومة مؤكدة", "[حقيقة موثقة أو قرار معتمد يمكن البناء عليه]"),
            ("معلومة منقولة", "[معلومة ذكرها طرف معني ولم تُتحقق بصورة مستقلة]"),
            ("فرضية للاختبار", "[تفسير أو اتجاه يحتاج إلى تجربة أو تحقق]"),
        ],
        [2200, content_width - 2200],
    )

    doc.add_paragraph("معايير النجاح المقترحة", style="Heading 2")
    add_bullet(doc, "[مؤشر قابل للقياس] — خط الأساس: [القيمة]، والهدف المقترح: [القيمة]، وطريقة القياس: [الطريقة].", bullet_id)
    add_bullet(doc, "[مؤشر جودة أو تجربة] — حد القبول المقترح: [الحد]، ومسؤول القياس: [المالك].", bullet_id)
    add_callout(
        doc,
        "تظل الأهداف والمؤشرات مقترحة إلى أن تُعتمد بيانات خط الأساس وطريقة القياس والمالك المسؤول.",
        content_width,
        label="تنبيه اعتماد",
        fill=MIST,
        accent=BLUE_DEEP,
    )

    doc.add_paragraph("القرار المطلوب", style="Heading 2")
    add_text_paragraph(
        doc,
        "[اكتب القرار الذي تحتاجه من العميل الآن، وسبب الحاجة إليه، وصاحب الصلاحية المطلوب اعتماده.]",
        style="INTAG Placeholder",
    )

    doc.add_paragraph("الأسئلة المفتوحة", style="Heading 2")
    add_bullet(doc, "[السؤال المؤثر في النطاق] — المالك: [الاسم] — موعد الإجابة: [التاريخ].", bullet_id)
    add_bullet(doc, "[السؤال المؤثر في القياس أو القرار] — المالك: [الاسم] — موعد الإجابة: [التاريخ].", bullet_id)

    add_page_break(doc)

    # الصفحة 3: النطاق والمنهج.
    p = doc.add_paragraph(style="INTAG Kicker")
    p.add_run("02 / النطاق والمنهج")
    doc.add_paragraph("النطاق والمنهج", style="Heading 1")
    add_text_paragraph(
        doc,
        "حوّل الهدف إلى مخرجات محددة ودليل قبول ومسؤول واضح. لا يبقى في النسخة النهائية إلا ما اعتمده أصحاب القرار.",
        style="Normal",
    )

    scope_widths = [1900, 2780, 2780, 1900]
    add_data_table(
        doc,
        ["مجال العمل", "المخرج المقترح", "دليل القبول", "الحالة"],
        [
            ("الاستراتيجية واستشارات الأعمال", "[مخرج استراتيجي محدد]", "[قرار أو وثيقة اعتماد]", "قيد المراجعة"),
            ("العلامة التجارية والإعلام", "[أصل بصري أو محتوى محدد]", "[معايير مراجعة وقائمة أصول]", "قيد المراجعة"),
            ("التقنية والمنتجات الرقمية", "[حل أو منتج رقمي محدد]", "[اختبار أو عرض أو تسليم موثق]", "قيد المراجعة"),
            ("التسويق والنمو", "[تجربة أو حملة محددة]", "[نتيجة مقاسة أو تعلم موثق]", "قيد المراجعة"),
        ],
        scope_widths,
        status_column=3,
    )

    doc.add_paragraph("منهج العمل", style="Heading 2")
    add_number(doc, "الاكتشاف والتحديد: مراجعة السياق والبيانات والقيود وتثبيت السؤال المطلوب حسمه.", number_id)
    add_number(doc, "التصميم والتجريب: إعداد اتجاه قابل للمراجعة واختباره ضمن حدود متفق عليها.", number_id)
    add_number(doc, "التحقق والتحسين: قياس الدليل ومراجعة الملاحظات وتصحيح الفرضيات.", number_id)
    add_number(doc, "التسليم والاعتماد: توثيق المخرجات والقرارات والملكية والخطوة التالية.", number_id)

    doc.add_paragraph("حدود النطاق", style="Heading 2")
    add_bullet(doc, "لا يدخل في النطاق إلا ما يرد صراحة في النسخة المعتمدة من هذا المقترح.", bullet_id)
    add_bullet(doc, "تظل التراخيص الخارجية والإنفاق الإعلامي والاستضافة والصيانة والتكاملات قيد التحديد ما لم تُذكر صراحة.", bullet_id)
    add_bullet(doc, "أي تغيير مؤثر في النطاق أو المدة أو التكلفة يحتاج إلى قرار تغيير موثق قبل التنفيذ.", bullet_id)

    doc.add_paragraph("التبعيات والمدخلات", style="Heading 2")
    add_bullet(doc, "من العميل: [بيانات أو أصول أو صلاحيات أو ملاحظات لازمة].", bullet_id)
    add_bullet(doc, "من إنتاغ: [مالك التنفيذ ومعايير المراجعة والتسليم].", bullet_id)
    add_bullet(doc, "من طرف ثالث: [ترخيص أو موافقة أو خدمة خارجية، إن وجدت].", bullet_id)

    add_page_break(doc)

    # الصفحة 4: الجدول والمسؤوليات.
    p = doc.add_paragraph(style="INTAG Kicker")
    p.add_run("03 / الجدول الزمني والمسؤوليات")
    doc.add_paragraph("الجدول الزمني والمسؤوليات", style="Heading 1")
    add_text_paragraph(
        doc,
        "تبدأ المدد بعد اكتمال المدخلات والاعتمادات المرتبطة بكل مرحلة، ولا تعد المدة التقديرية وعدًا نهائيًا قبل اعتماد النطاق.",
        style="Normal",
    )

    timeline_widths = [1800, 2100, 3300, 2160]
    add_data_table(
        doc,
        ["المرحلة", "المدة التقديرية", "المخرج الرئيس", "المالك"],
        [
            ("الاكتشاف والتحديد", "[عدد الأيام أو الأسابيع]", "ملخص الفهم والنطاق الأولي", "[اسم المالك]"),
            ("التصميم والتجريب", "[عدد الأيام أو الأسابيع]", "نسخة قابلة للمراجعة والاختبار", "[اسم المالك]"),
            ("التحقق والتحسين", "[عدد الأيام أو الأسابيع]", "نتائج القياس والتعديلات المتفق عليها", "[اسم المالك]"),
            ("التسليم والاعتماد", "[عدد الأيام أو الأسابيع]", "حزمة التسليم وسجل القرارات", "[اسم المالك]"),
        ],
        timeline_widths,
    )

    doc.add_paragraph("مسؤوليات إنتاغ", style="Heading 2")
    add_bullet(doc, "إدارة العمل المتفق عليه، وتوضيح الافتراضات والقيود، وإبقاء القرارات قابلة للتتبع.", bullet_id)
    add_bullet(doc, "تقديم المخرجات للمراجعة وفق النطاق، وتسجيل ما يحتاج إلى اعتماد أو تدخل من طرف آخر.", bullet_id)

    doc.add_paragraph("مسؤوليات العميل", style="Heading 2")
    add_bullet(doc, "توفير المدخلات والصلاحيات والأصول والملاحظات في الأوقات المتفق عليها.", bullet_id)
    add_bullet(doc, "تسمية أصحاب القرار واعتماد النطاق والمخرجات والتغييرات في صورة مكتوبة.", bullet_id)

    add_callout(
        doc,
        "إذا تأخر اعتماد أو مدخل أساسي، تُراجع المدة والتبعيات والأثر التجاري قبل تثبيت أي موعد جديد.",
        content_width,
        label="قاعدة الجدول الزمني",
        fill=MINERAL,
        accent=GREEN_DEEP,
    )

    doc.add_paragraph("نقاط المراجعة", style="Heading 2")
    add_data_table(
        doc,
        ["النقطة", "القرار المطلوب", "المعتمد"],
        [
            ("بعد الاكتشاف", "تثبيت الهدف والنطاق والقياس", "[الاسم]"),
            ("بعد النسخة الأولية", "اعتماد اتجاه التحسين أو طلب تعديل محدد", "[الاسم]"),
            ("قبل التسليم", "اعتماد المخرجات والملكية والخطوة التالية", "[الاسم]"),
        ],
        [2200, 4500, content_width - 6700],
    )

    add_page_break(doc)

    # الصفحة 5: الاستثمار والاعتماد.
    p = doc.add_paragraph(style="INTAG Kicker")
    p.add_run("04 / الاستثمار والاعتماد")
    doc.add_paragraph("الاستثمار والاعتماد", style="Heading 1")
    add_text_paragraph(
        doc,
        "استخدم هذا القسم لتسجيل أساس التقدير، لا لإدخال سعر نهائي قبل اكتمال المراجعة التجارية والفنية.",
        style="Normal",
    )

    investment_widths = [2460, 2880, 2260, 1760]
    add_data_table(
        doc,
        ["البند", "أساس التقدير", "القيمة المقترحة", "الحالة"],
        [
            ("أتعاب العمل", "[النطاق والجهد والمدة]", "[قيمة تُحدد بعد المراجعة]", "قيد المراجعة"),
            ("خدمات أو تراخيص خارجية", "[المورد والاستخدام المتوقع]", "[تقدير منفصل]", "قيد المراجعة"),
            ("الإنتاج أو الإنفاق الإعلامي", "[الخطة وحجم التنفيذ]", "[تقدير منفصل]", "قيد المراجعة"),
            ("الضرائب والمصروفات", "[وفق الاتفاق والقواعد المطبقة]", "[تُحدد تعاقديًا]", "قيد المراجعة"),
        ],
        investment_widths,
        status_column=3,
    )

    doc.add_paragraph("أسس تجارية تحتاج إلى اعتماد", style="Heading 2")
    add_bullet(doc, "صلاحية العرض: [المدة المقترحة بعد الاعتماد].", bullet_id)
    add_bullet(doc, "جدول الدفعات: [يحدده المسؤول التجاري ويعتمد كتابيًا].", bullet_id)
    add_bullet(doc, "شروط التغيير والإلغاء: [تراجع وتوثق في الاتفاق النهائي].", bullet_id)

    add_callout(
        doc,
        "لا يضمن هذا المقترح مبيعات أو نموًا أو دقة أو موعد تسليم أو جاهزية تشغيلية. يصبح النطاق والتكلفة والمدة ملزمة فقط بعد اعتمادها من أصحاب الصلاحية في وثيقة نهائية.",
        content_width,
        label="حدود الالتزام",
        fill=MIST,
        accent=BLUE_DEEP,
    )

    doc.add_paragraph("الاعتماد", style="Heading 2")
    approval = add_data_table(
        doc,
        ["اعتماد العميل", "اعتماد إنتاغ"],
        [
            (
                "الاسم: [اسم المعتمد]\nالصفة: [الصفة]\nالتوقيع: [قيد الاستكمال]\nالتاريخ: [التاريخ]",
                "الاسم: [اسم المعتمد]\nالصفة: [الصفة]\nالتوقيع: [قيد الاستكمال]\nالتاريخ: [التاريخ]",
            )
        ],
        [half, content_width - half],
    )
    for cell in approval.rows[1].cells:
        cell.paragraphs[0].paragraph_format.line_spacing = 1.25

    finalize_document(doc, PROPOSAL_OUT)


def build_letterhead() -> None:
    doc = Document()
    section = doc.sections[0]
    content_width = configure_a4_section(section, top=1.04, bottom=0.72)
    configure_styles(doc, preset="standard_business_brief")
    set_header_footer(
        section,
        document_label="نموذج مراسلات إنتاغ",
        status_label="الشعار 02 معتمد، وبقية القرارات قيد المراجعة",
        content_width=content_width,
        logo=True,
    )
    add_doc_properties(
        doc,
        title="نموذج المراسلات الرسمية لإنتاغ - الإصدار 2.0",
        subject="ورق رسمي عربي عملي وهادئ قابل للتحرير",
    )

    p = doc.add_paragraph(style="INTAG Kicker")
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(11)
    p.add_run("مراسلات رسمية")

    meta = doc.add_paragraph(style="INTAG Meta")
    set_paragraph_rtl(meta)
    r = meta.add_run("التاريخ: [التاريخ]  ·  المرجع: [رقم المرجع]  ·  الإصدار: 2.0")
    set_run_font(r, size=9, color=SLATE, bold=True)
    meta.paragraph_format.space_after = Pt(14)

    for line, bold in (
        ("[اسم المستلم]", True),
        ("[الصفة الوظيفية]", False),
        ("[اسم الجهة]", False),
        ("[العنوان]", False),
        ("[المدينة والدولة]", False),
    ):
        p = doc.add_paragraph()
        set_paragraph_rtl(p)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        set_run_font(r, size=10, color=INK, bold=bold)

    add_spacer(doc, 7)
    add_callout(
        doc,
        "[اكتب عنوانًا مختصرًا ومحددًا للرسالة]",
        content_width,
        label="الموضوع",
        fill=MINERAL,
        accent=GREEN_DEEP,
    )

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("تحية طيبة وبعد،")
    set_run_font(r, size=11, color=INK)

    for text in (
        "[اكتب الفقرة الافتتاحية بوضوح، واذكر الغرض من الرسالة والسياق الضروري فقط.]",
        "[اعرض المعلومات أو القرارات المطلوبة في فقرات قصيرة، وتحقق من الأسماء والتواريخ والمالكين قبل الإرسال.]",
        "[اختم بخطوة تالية واضحة، وحدد من سيتولى المتابعة ومتى، من دون إضافة التزام غير معتمد.]",
    ):
        p = doc.add_paragraph(style="INTAG Placeholder")
        set_paragraph_rtl(p)
        p.paragraph_format.space_after = Pt(11)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=SLATE)

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("وتفضلوا بقبول خالص التحية،")
    set_run_font(r, size=10.5, color=INK)

    for line, bold, color in (
        ("[الاسم الكامل]", True, MIDNIGHT),
        ("[المسمى الوظيفي]", False, SLATE),
        ("إنتاغ للحلول الرقمية", True, BLUE_DEEP),
        ("[البريد الإلكتروني]  ·  [رقم الهاتف]  ·  [الموقع الإلكتروني]", False, SLATE),
    ):
        p = doc.add_paragraph()
        set_paragraph_rtl(p)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run_font(r, size=9.5 if not bold else 10, color=color, bold=bold)

    add_spacer(doc, 22)
    add_callout(
        doc,
        "استبدل جميع الحقول المحاطة بأقواس مربعة، ثم راجع الحقائق والتواريخ والصلاحيات والالتزامات قبل الإرسال.",
        content_width,
        label="ملاحظة القالب",
        fill=MIST,
        accent=BLUE_DEEP,
    )

    finalize_document(doc, LETTERHEAD_OUT)


if __name__ == "__main__":
    if not LOGO.exists():
        raise FileNotFoundError(LOGO)
    verify_logo_geometry()
    build_proposal()
    build_letterhead()
    print(PROPOSAL_OUT)
    print(LETTERHEAD_OUT)
